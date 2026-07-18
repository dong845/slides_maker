#!/usr/bin/env python3
"""ingest — turn a NON-PDF source (Office doc, spreadsheet, image, audio, video) into something the
content-planner can read PRECISELY, or say honestly when it can't. Companion to extract_pdf.py (PDFs)
and extract_deck.py (existing .pptx).

  probe   <file>                  Detect the type and print the recommended ingest route.
  doctext <file.docx> [out.txt]   Word .docx text + tables DIRECTLY (python-docx), heading-tagged —
                                  the HIGHEST-fidelity path for a short/word-primary Word doc (exact
                                  text, not a vision guess). A long .docx: use `office` -> PDF instead.
  sheet   <file.xlsx> [out.csv]   Spreadsheet cells -> CSV via openpyxl (EXACT rows). `office`->PDF is
                                  layout-only and DROPS/truncates the data, so use `sheet` for numbers.
  office  <file> [outdir]         Office doc (.docx/.doc/.odt/...) -> PDF via LibreOffice, so the full
                                  extract_pdf pipeline (map/text/figures/headings, incl. long-source
                                  mode) applies. Best when LAYOUT / figures matter. (.pptx -> use
                                  extract_deck.py natively; office->PDF is a lossy fallback.)
  frames  <video> <outdir> [--every SEC]   Sample keyframes from a video (ffmpeg) for VISUAL reading
                                  (slides shown, UI, scenes), capped at 60 frames. Prints duration + a
                                  ⚠: the SPOKEN track needs a transcript — there is NO speech-to-text.

Fidelity floor (the planner + critic enforce it): `doctext`/`sheet`/`office` text is EXACT (traceable,
high-fidelity). An IMAGE / VIDEO FRAME is read by the model's own vision (no OCR installed) — good to
understand & plan, but a number/quote taken off pixels is a claim-ledger `verified? = N` row until
confirmed against real data/text. AUDIO and a video's narration are unavailable without a supplied
transcript (.srt/.vtt/.txt) or an installed STT tool — never invent what you couldn't hear/read.
"""
import sys
import os
import shutil
import tempfile
import subprocess

AUDIO_EXT = (".mp3", ".m4a", ".wav", ".aac", ".flac", ".ogg", ".opus", ".wma")


def _die(msg):
    raise SystemExit(f"error: {msg}")


def probe(path):
    if not os.path.exists(path):
        _die(f"no such file: {path}  (a Google Doc / Slides / Notion / URL isn't a local file — "
             "export or download it first, or paste the text)")
    ext = os.path.splitext(path)[1].lower()
    audio = "Audio — NO speech-to-text installed; supply a transcript (.srt/.vtt/.txt) for the spoken content, else it's a gap (verified? = N). Never invent narration."
    vid = "Video — `ingest.py frames` for VISUALS; supply a transcript for the SPOKEN content (no STT here)."
    img = "Image — read with the model's vision (no OCR). Place the original as a figure freely; a number/quote you TYPE off it is verified? = N until confirmed."
    routes = {
        ".pdf": "PDF — use extract_pdf.py (map/text/figures/headings); ingest.py not needed.",
        ".docx": "Word — `ingest.py doctext` for exact text+tables (short docs), or `office` -> PDF for layout / a long doc.",
        ".doc": "Word (legacy) — `ingest.py office` -> PDF (python-docx can't read .doc).",
        ".pptx": "PowerPoint — use extract_deck.py (native, keeps figures); `office` -> PDF only for a flat read.",
        ".ppt": "PowerPoint (legacy) — `ingest.py office` -> PDF.",
        ".xlsx": "Excel — `ingest.py sheet` (exact rows via openpyxl). `office`->PDF is layout-only and DROPS data — don't use it for numbers.",
        ".xlsm": "Excel — `ingest.py sheet` (exact rows via openpyxl).",
        ".xls": "Excel (legacy) — re-save as .xlsx then `ingest.py sheet`, or `office` -> PDF for a printed view.",
        ".csv": "CSV — read directly (already plain rows).",
        ".odt": "OpenDocument text — `ingest.py office` -> PDF.",
        ".md": "Markdown — read as text directly; convert to PDF only if you need paginated triage.",
        ".txt": "Plain text — read directly.",
        ".epub": "EPUB — extract_pdf.py opens it (map/text/headings) like a PDF.",
        ".png": img, ".jpg": img, ".jpeg": img, ".gif": img, ".webp": img, ".bmp": img, ".tiff": img,
        ".mp4": vid, ".mov": vid, ".mkv": vid, ".webm": vid, ".avi": vid, ".m4v": vid,
        ".srt": "Transcript/captions — read as text; the precise spoken content of a video.",
        ".vtt": "Transcript/captions — read as text; the precise spoken content of a video.",
    }
    for a in AUDIO_EXT:
        routes[a] = audio
    print(f"{path}  [{ext or 'no-ext'}]")
    print("  route:", routes.get(ext, "unknown type — if it's a document, convert to PDF and use extract_pdf.py."))
    return 0


def doctext(path, out=None):
    if not os.path.exists(path):
        _die(f"no such file: {path}")
    if not path.lower().endswith(".docx"):
        _die(f"doctext reads .docx only (got {os.path.splitext(path)[1] or 'no-ext'}); "
             "for .doc/.odt use `office` -> PDF.")
    try:
        import docx  # python-docx
    except ImportError:
        _die("python-docx not installed (`pip install python-docx`); or use `office` -> PDF instead.")
    try:
        d = docx.Document(path)
    except Exception as e:
        _die(f"can't open {path!r} as a .docx ({e.__class__.__name__}: {e}) — is it a real, "
             "non-corrupt Word file? (a renamed .doc won't work — use `office`.)")
    lines = []
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = (p.style.name if p.style else "") or ""
        lines.append(f"## {t}" if style.lower().startswith("heading") else t)
    for i, tbl in enumerate(d.tables, 1):
        lines.append(f"\n[TABLE {i}]")
        for row in tbl.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    # textboxes/callouts — python-docx paragraphs can't see them, but the XML can; a load-bearing
    # number in a callout must not vanish silently
    try:
        tb_texts = [t.text for t in d.element.body.xpath(".//w:txbxContent//w:t") if t.text]
        if tb_texts:
            lines.append("\n[TEXTBOXES]")
            lines.append(" ".join(tb_texts))
    except Exception:
        pass
    # section headers/footers (skip empties)
    hf = []
    for sec in d.sections:
        for part, tag in ((sec.header, "HEADER"), (sec.footer, "FOOTER")):
            for p in part.paragraphs:
                t = p.text.strip()
                if t:
                    hf.append(f"[{tag}] {t}")
    if hf:
        lines.append("")
        lines.extend(dict.fromkeys(hf))          # dedupe repeated per-section chrome, keep order
    text = "\n".join(lines)
    # honest limits: footnotes/endnotes + SDT content are still invisible to python-docx
    try:
        import zipfile
        with zipfile.ZipFile(path) as z:
            if any(n in z.namelist() for n in ("word/footnotes.xml", "word/endnotes.xml")):
                print("⚠ this .docx has footnotes/endnotes — doctext does NOT extract them; "
                      "use `office` -> PDF if they carry content.")
    except Exception:
        pass
    if not text.strip():
        print("⚠ no extractable text (empty, or text lives in images/scans python-docx can't see). "
              "If it's SCANNED, `office` -> PDF gives image PAGES you read with VISION — there is no "
              "OCR here, so those are `verified? = N` (not recoverable as exact text). Ask for a "
              "text/Word original if the exact wording is load-bearing.")
    if out:
        with open(out, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"wrote {path} -> {out} ({len(text.split()):,} words, {len(d.tables)} tables)")
    else:
        print(text)
    return 0


def sheet(path, out=None):
    if not os.path.exists(path):
        _die(f"no such file: {path}")
    if not path.lower().endswith((".xlsx", ".xlsm")):
        _die(f"sheet reads .xlsx/.xlsm (openpyxl); got {os.path.splitext(path)[1] or 'no-ext'} — "
             "for .xls/.ods, re-save as .xlsx or use `office` -> PDF.")
    try:
        import openpyxl
    except ImportError:
        _die("openpyxl not installed (`pip install openpyxl`).")
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception as e:
        _die(f"can't open {path!r} as a spreadsheet ({e.__class__.__name__}: {e}).")
    import csv
    import io
    import datetime as _dt
    buf = io.StringIO()
    w = csv.writer(buf)
    rows = 0
    blanks = 0                                    # None cells inside emitted rows — formula suspects

    def _cell(c):
        if isinstance(c, _dt.datetime) and c.time() == _dt.time(0, 0):
            return c.date().isoformat()           # pure dates, not spurious midnight timestamps
        return c

    for ws in wb.worksheets:
        buf.write(f"\n===== SHEET: {ws.title} =====\n")
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                blanks += sum(1 for c in row if c is None)
                w.writerow(["" if c is None else _cell(c) for c in row])
                rows += 1
    n_sheets = len(wb.worksheets)
    wb.close()
    text = buf.getvalue()
    if rows == 0:
        print("⚠ no cell data found (empty workbook, or values are formulas not yet computed).")
    elif blanks:
        # data_only=True returns None for a formula cell whose file carries no cached value (any
        # programmatically-written workbook never re-saved by a spreadsheet app) — the "exact rows"
        # promise silently breaks unless we check
        try:
            import openpyxl as _px
            wb2 = _px.load_workbook(path, read_only=True, data_only=False)
            n_formula = sum(1 for ws in wb2.worksheets for row in ws.iter_rows()
                            for c in row if c.data_type == "f")
            wb2.close()
            if n_formula:
                print(f"⚠ {n_formula} formula cell(s) have no cached value — shown as EMPTY in the "
                      "CSV. Open + re-save the file in Excel/LibreOffice (computes the cache), or "
                      "get a CSV export, before treating these columns as absent.")
        except Exception:
            pass
    if out:
        with open(out, "w", encoding="utf-8", newline="") as f:
            f.write(text)
        print(f"wrote {path} -> {out} ({n_sheets} sheet(s), {rows:,} non-empty rows)")
    else:
        print(text)
    return 0


def office(path, outdir="."):
    if not os.path.exists(path):
        _die(f"no such file: {path}")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        _die("LibreOffice (soffice) not found — install it, or export the doc to PDF yourself.")
    os.makedirs(outdir, exist_ok=True)
    profile = tempfile.mkdtemp(prefix="lo_profile_")   # OUTSIDE outdir → never litters the deliverable
    workdir = tempfile.mkdtemp(prefix="lo_out_")       # fresh dir → a stale same-name PDF in outdir
    #                                                    can never masquerade as a fresh conversion
    cmd = [soffice, "--headless", f"-env:UserInstallation=file://{profile}",
           "--convert-to", "pdf", "--outdir", workdir, path]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
    except subprocess.TimeoutExpired:
        shutil.rmtree(workdir, ignore_errors=True)
        _die("LibreOffice conversion timed out (180s).")
    finally:
        shutil.rmtree(profile, ignore_errors=True)
    name = os.path.splitext(os.path.basename(path))[0] + ".pdf"
    tmp_pdf = os.path.join(workdir, name)
    if r.returncode != 0 or not os.path.exists(tmp_pdf):
        shutil.rmtree(workdir, ignore_errors=True)
        _die(f"conversion FAILED (rc={r.returncode}). soffice said: "
             f"{(r.stderr or r.stdout).strip()[:300]}")
    out_pdf = os.path.join(outdir, name)
    shutil.move(tmp_pdf, out_pdf)
    shutil.rmtree(workdir, ignore_errors=True)
    print(f"converted {path} -> {out_pdf}  (now run: extract_pdf.py map {out_pdf})")
    return 0


def _duration(video):
    ffprobe = shutil.which("ffprobe")
    if not ffprobe:
        return None
    try:
        r = subprocess.run([ffprobe, "-v", "error", "-show_entries", "format=duration",
                            "-of", "default=nw=1:nk=1", video], capture_output=True, text=True, timeout=30)
        return float(r.stdout.strip())
    except Exception:
        return None


def _has_video_stream(video):
    ffprobe = shutil.which("ffprobe")
    if not ffprobe:
        return True   # can't tell → let ffmpeg try
    try:
        r = subprocess.run([ffprobe, "-v", "error", "-select_streams", "v", "-show_entries",
                            "stream=codec_type", "-of", "csv=p=0", video],
                           capture_output=True, text=True, timeout=30)
        return "video" in r.stdout
    except Exception:
        return True


def frames(video, outdir, every=5.0, cap=60):
    if not os.path.exists(video):
        _die(f"no such file: {video}")
    if outdir.startswith("-"):
        _die(f"refusing outdir {outdir!r} (starts with '-') — did you omit the outdir argument?")
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        _die("ffmpeg not found — install it to sample video frames.")
    if every <= 0:
        _die("--every must be > 0 seconds")
    if not _has_video_stream(video):
        _die(f"{video!r} has no video track (audio-only?). There is NO speech-to-text here — supply "
             "a transcript / captions (.srt/.vtt/.txt) for the spoken content; never invent it.")
    os.makedirs(outdir, exist_ok=True)
    dur = _duration(video)
    if dur and int(dur / every) + 1 > cap:            # keep the model's reading load bounded
        every = dur / cap
        print(f"⚠ clamping to --every {every:.1f}s to stay under the fixed {cap}-frame cap for a "
              f"{dur:.0f}s video (a larger --every changes the spacing; to see a region closely, "
              "extract a sub-clip with ffmpeg and sample that).")
    pattern = os.path.join(outdir, "frame_%04d.png")
    cmd = [ffmpeg, "-y", "-loglevel", "error", "-i", video, "-vf", f"fps=1/{every}",
           "-frames:v", str(cap), pattern]        # -frames:v is a hard ceiling even if duration unknown
    try:
        subprocess.run(cmd, capture_output=True, text=True, timeout=600, check=True)
    except subprocess.CalledProcessError as e:
        _die(f"ffmpeg failed ({(e.stderr or '').strip()[:300]}) — is {video!r} a readable video?")
    except subprocess.TimeoutExpired:
        _die("frame extraction timed out (600s) — sample a shorter clip or a larger --every.")
    n = len([f for f in os.listdir(outdir) if f.startswith("frame_") and f.endswith(".png")])
    dtxt = f"{dur:.0f}s" if dur else "unknown duration"
    print(f"extracted {n} frames (1 per {every:g}s, {dtxt}) -> {outdir}/frame_*.png — read them with vision.")
    print("⚠ VISUAL only: the SPOKEN narration is NOT captured (no speech-to-text here). For the "
          "audio content, supply a transcript / captions (.srt/.vtt/.txt). Don't invent narration.")
    return 0


def _main(argv):
    if len(argv) < 2:
        print(__doc__)
        return 1
    cmd, a = argv[1], argv[2:]
    try:
        if cmd == "probe":
            if not a:
                print("usage: ingest.py probe <file>"); return 1
            return probe(a[0])
        if cmd == "doctext":
            if not a:
                print("usage: ingest.py doctext <file.docx> [out.txt]"); return 1
            return doctext(a[0], a[1] if len(a) > 1 else None)
        if cmd == "sheet":
            if not a:
                print("usage: ingest.py sheet <file.xlsx> [out.csv]"); return 1
            return sheet(a[0], a[1] if len(a) > 1 else None)
        if cmd == "office":
            if not a:
                print("usage: ingest.py office <file> [outdir]"); return 1
            return office(a[0], a[1] if len(a) > 1 else ".")
        if cmd == "frames":
            every = 5.0
            if "--every" in a:                    # strip flag+value BEFORE reading positionals
                i = a.index("--every")
                try:
                    every = float(a[i + 1])
                except (IndexError, ValueError):
                    print("error: --every needs a number of seconds"); return 1
                a = a[:i] + a[i + 2:]
            if len(a) < 2:
                print("usage: ingest.py frames <video> <outdir> [--every SEC]"); return 1
            return frames(a[0], a[1], every=every)
        print(__doc__)
        return 1
    except SystemExit:
        raise
    except Exception as e:                        # never dump a raw traceback at an agent
        print(f"error: {e.__class__.__name__}: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(_main(sys.argv))
